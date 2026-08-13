from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK

OUT = 'MmaamulPro_Client_User_Manual.docx'

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), fill); tcPr.append(shd)
def border_bottom(p, color='1F4E79'):
    pPr=p._p.get_or_add_pPr(); b=OxmlElement('w:pBdr'); bot=OxmlElement('w:bottom'); bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'12'); bot.set(qn('w:space'),'6'); bot.set(qn('w:color'),color); b.append(bot); pPr.append(b)
def page_field(p):
    r=p.add_run(); fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'PAGE'); r._r.append(fld)
def set_cell(cell, text, bold=False, color=None):
    cell.text=''; p=cell.paragraphs[0]; p.paragraph_format.space_after=Pt(0); r=p.add_run(str(text)); r.bold=bold
    if color: r.font.color.rgb=RGBColor.from_string(color)
    cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
def table(doc, headers, rows, widths=None):
    t=doc.add_table(rows=1, cols=len(headers)); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
    trPr = t.rows[0]._tr.get_or_add_trPr(); tblHeader = OxmlElement('w:tblHeader'); tblHeader.set(qn('w:val'), 'true'); trPr.append(tblHeader)
    for i,h in enumerate(headers):
        set_cell(t.rows[0].cells[i],h,True,'FFFFFF'); shade(t.rows[0].cells[i],'1F4E79')
        if widths: t.rows[0].cells[i].width=Inches(widths[i])
    for row in rows:
        cells=t.add_row().cells
        for i,val in enumerate(row):
            set_cell(cells[i],val)
            if widths: cells[i].width=Inches(widths[i])
    doc.add_paragraph('')
    return t
def callout(doc, label, text, color='EAF2F8'):
    t=doc.add_table(rows=1, cols=1); t.autofit=False; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    trPr = t.rows[0]._tr.get_or_add_trPr(); tblHeader = OxmlElement('w:tblHeader'); tblHeader.set(qn('w:val'), 'true'); trPr.append(tblHeader)
    c=t.cell(0,0); shade(c,color); p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(2); r=p.add_run(label+'  '); r.bold=True; r.font.color.rgb=RGBColor(31,78,121); p.add_run(text)
    doc.add_paragraph('')
def add_bullets(doc, items):
    for x in items: doc.add_paragraph(x, style='List Bullet')
def steps(doc, items):
    for x in items: doc.add_paragraph(x, style='List Number')

d=Document(); sec=d.sections[0]
sec.top_margin=sec.bottom_margin=Inches(0.85); sec.left_margin=sec.right_margin=Inches(0.85); sec.header_distance=Inches(.35); sec.footer_distance=Inches(.35)
styles=d.styles
styles['Normal'].font.name='Calibri'; styles['Normal'].font.size=Pt(10.5); styles['Normal'].paragraph_format.space_after=Pt(5); styles['Normal'].paragraph_format.line_spacing=1.1
for name,size,color,before,after in [('Title',30,'1F4E79',0,12),('Heading 1',17,'1F4E79',16,8),('Heading 2',13,'2E74B5',12,5),('Heading 3',11,'1F4E79',8,4)]:
    s=styles[name]; s.font.name='Calibri'; s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=RGBColor.from_string(color); s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after); s.paragraph_format.keep_with_next=True
for n in ['List Bullet','List Number']:
    styles[n].font.name='Calibri'; styles[n].font.size=Pt(10.5); styles[n].paragraph_format.space_after=Pt(3)
header=sec.header.paragraphs[0]; header.text='MAAMULPRO  |  CLIENT USER MANUAL'; header.style='Caption'; border_bottom(header); header.runs[0].font.color.rgb=RGBColor(31,78,121)
footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.RIGHT; footer.add_run('MmaamulPro Client User Manual  |  Page '); page_field(footer)

# Cover
d.add_paragraph('')
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(110); r=p.add_run('MmaamulPro'); r.font.size=Pt(34); r.bold=True; r.font.color.rgb=RGBColor(31,78,121)
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('Client User Manual'); r.font.size=Pt(25); r.bold=True; r.font.color.rgb=RGBColor(46,116,181)
p=d.add_paragraph('Construction | Real Estate | Materials | Finance | People'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.runs[0].font.size=Pt(13); p.runs[0].font.color.rgb=RGBColor(89,89,89)
d.add_paragraph('')
callout(d,'Purpose','A practical, client-facing guide to operating MmaamulPro from daily work through management review.','EAF2F8')
p=d.add_paragraph('Version: Client edition | Prepared: 13 August 2026'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
d.add_page_break()

d.add_heading('Table of Contents',0)
toc=[
('1. Introduction to MmaamulPro',),('2. Getting Started',),('3. Navigation and Dashboard',),('4. User Roles and Permissions',),('5. Core Operations: Staff, Financials and Payroll',),('6. Construction Management',),('7. Real Estate Management',),('8. Materials Management',),('9. Reports, Analytics and Tracking',),('10. Administration and Settings',),('11. Complete Business Workflows',),('12. Common Tasks / Quick Guides',),('13. Troubleshooting and Common Mistakes',),('14. Best Practices',),('15. Glossary',)]
table(d,['Section'],toc,[6.3])
callout(d,'How to use this manual','Use the numbered sections for full procedures. Use the Quick Guides for common repeat tasks. Menu visibility depends on the permissions assigned to your account.')
d.add_page_break()

def h1(title): d.add_heading(title,1)
def h2(title): d.add_heading(title,2)
def h3(title): d.add_heading(title,3)
def purpose(text): d.add_paragraph(text)
def page_guide(when, steps_list, track, related='', notes=None):
    h3('When to use it'); d.add_paragraph(when)
    h3('How to use it'); steps(d,steps_list)
    h3('Where to track the result'); d.add_paragraph(track)
    if related: h3('Related modules'); d.add_paragraph(related)
    if notes: callout(d,'Important',notes,'FFF4E5')

h1('1. Introduction to MmaamulPro')
purpose('MmaamulPro is a unified business-management system for organisations that operate construction projects, real estate portfolios and material sales. It combines operational records with financial tracking so that work completed in a business module can be reviewed through the appropriate list, detail page, report and, where applicable, the financial ledger.')
table(d,['Workspace','What it manages','Typical users'],[
('Core','Staff, financial transactions, accounts, journals, payroll, reports and company settings.','Owners, finance, HR and managers'),('Construction','Projects, tasks, site costs, labour, contracts and construction inventory.','Construction managers, supervisors and storekeepers'),('Real Estate','Properties, clients, deals, tenancies, rental contracts and rent collections.','Property and rental teams'),('Materials','Products, suppliers, purchase orders, sales invoices, customers and deliveries.','Procurement, sales and delivery teams')],[1.25,3.55,1.5])
callout(d,'Important','This manual describes the behaviour present in the installed system. Your company may not see every menu because access is controlled by role and individual permission.')

h1('2. Getting Started')
h2('2.1 Sign in and recover access')
page_guide('Use Sign in at the start of each work session. Use Forgot password only when you cannot access your account.',[
'Open the MmaamulPro sign-in page and enter your company account email and password.',
'Select Sign in. The system opens the dashboard or the home page permitted for your role.',
'If you cannot sign in, select Forgot password and follow the recovery instructions sent for your account.',
'If the account is locked, contact a company administrator; they can activate the staff account or reset its password from Staff Management.'
],'Your profile/session controls are in the top navigation; account status is managed in Staff Management.','Staff Management and Roles & Permissions.','Passwords set by an administrator require at least six characters. A role or direct-permission change ends the affected user session; the user must sign in again.')
h2('2.2 First-time company setup')
page_guide('Use Settings before operational work begins, and create staff records before assigning work, payroll or user access.',[
'Open Settings and complete the company information and operational preferences available to your role.',
'Open Staff Management and create the employees or workers who will be selected in projects, tasks, contracts or payroll.',
'Open Roles & Permissions to review standard roles or create a role that matches your internal separation of duties.',
'Create the operational master records needed by your workspace: properties, material products, suppliers, clients or projects.',
'Confirm that finance users have reviewed accounts and mappings before expecting automatic financial posting.'
],'Setup records remain available in their owning module; settings and access changes can be reviewed later by authorised administrators.','Settings, Staff, Roles, Financial Categories and Accounts.')

h1('3. System Navigation and Dashboard')
h2('3.1 Navigation principles')
purpose('The left navigation groups pages by workspace. Expand a group to open its pages; the active page is highlighted. The header provides account/session controls and notifications. Lists typically offer search, filters, row actions and a Create/Add action when you have the required permission.')
table(d,['Control','Use'],[('Search','Find records by name, reference or matching text.'),('Filters','Limit the list by status, project, period, department or another displayed criterion.'),('Status pill','Shows the current operating state; available actions change with the state.'),('Add / Create','Opens a new record form when you have create permission.'),('Edit','Changes a record only when its status permits editing and you have update permission.'),('View / row click','Opens supporting details or an item-specific page.'),('Transitions','Buttons such as Submit, Approve, Receive or Mark paid move a record through a controlled workflow.')],[1.35,5.05])
h2('3.2 Executive Dashboard and Analytics')
page_guide('Use Dashboard for a cross-business summary and Analytics for trend-oriented management review.',[
'Open Dashboard from the main menu.',
'Review the available KPI cards, operational summaries and recent information for your permitted workspaces.',
'Open Analytics when you need a broader performance view; apply its workspace/date controls where shown.',
'Drill into the underlying workspace lists or reports to investigate a figure before taking action.'
],'Use workspace reports, financial reports and list pages to trace the source records behind a result.','All business workspaces, Reports and Financials.')

h1('4. User Roles and Permissions')
purpose('MmaamulPro uses roles and individual permissions. A role bundles access; a direct permission can explicitly allow or deny an action for one user. A user must have the workspace and page permission to see or use a protected page.')
table(d,['Role family','Typical scope'],[('Executive roles','Company Owner, General Manager, Admin and Manager can access broad cross-workspace information, subject to the configured template.'),('Construction roles','Construction Manager, Site Engineer, Project Supervisor, Storekeeper, Procurement Officer and Manpower Supervisor.'),('Real-estate roles','Real Estate Manager, Sales Agent, Rental Officer and Property Supervisor.'),('Materials roles','Material Manager, Sales Staff, Inventory Officer, Supplier Officer and Delivery Officer.'),('Staff','No workspace access until a role or direct permissions are assigned.')],[2.0,4.4])
h2('4.1 Roles & Permissions')
page_guide('Use this page when an authorised administrator needs to control what people can view or change.',[
'Open Roles & Permissions.',
'To create a role, enter a unique role key, a clear name and description, choose active status, and select the permissions it should contain.',
'To edit a role, update its name, description, active status or permission selection and save.',
'Open a user’s access area to assign one or more active roles, or apply an explicit Allow/Deny permission with a reason.',
'Ask the affected user to sign in again after saving.'
],'Review roles in the roles list and user access on the selected staff/user record.','Staff Management.','System roles cannot be deleted. Deleting a custom role removes its assignments and disables it. Use direct permissions sparingly so access stays understandable.')

h1('5. Core Operations: Staff, Financials and Payroll')
h2('5.1 Staff Management')
page_guide('Use Staff Management to maintain employees, workers, project assignments and login accounts.',[
'Open Staff and select Add staff member.',
'Enter first name and last name; add photo, phone, position, department, employment status, salary, hire date and notes as appropriate.',
'For Construction staff, choose an assigned project if applicable.',
'To create a login immediately, select Create a login account and enter email, role and a temporary password.',
'Save. Select a staff row later to view details, account state and recent activity.',
'From the staff detail window, administrators can create an account later, activate/deactivate it, change email or role, or reset the password.'
],'The staff table shows department, position, workforce status, account availability and salary. Use search plus department/status filters.','Projects, Tasks, Manpower, Workforce Contracts and Payroll.','Do not delete a staff member until you confirm that historical records no longer need the person as a selectable operational record. Do not deactivate your own account.')
h2('5.2 Financial Transactions and Categories')
page_guide('Use Financials to record income and expenses that do not originate from a specialised workflow, or to review unified transaction history.',[
'Create financial categories first if a reusable classification is needed: enter Name, optional Code, Color and Description.',
'Open Financials and select the option to create a transaction.',
'Choose income or expense, enter amount, date, description and category, plus linked operational information where available.',
'Set the status accurately. A CLEARED transaction can be posted into the accounting journal when the required mapping is available.',
'Save and use the transaction list filters and summary to review the record.'
],'The Financials list and transaction detail are the primary follow-up locations; posted entries are also visible in Journal Entries and financial reports.','Categories, Accounts, Journal Entries, Projects, Deals, Rentals, Purchases, Sales and Payroll.','Editing or deleting a posted transaction triggers reversal handling for the earlier journal batch and re-posting only when the revised record remains cleared.')
h2('5.3 Accounts and Journal Entries')
page_guide('Use Accounts to maintain your chart of accounts, and Journal Entries to inspect or post balanced accounting batches.',[
'In Accounts, create an account with a unique code, name, type (Income, Expense, Asset, Liability or Equity), optional parent and balance rules.',
'Keep an account active when it is still available for posting; review relationships before changing or retiring it.',
'In Journal Entries, create a manual batch only when you need an accounting adjustment outside automated workflows.',
'Enter the batch date, memo and at least two debit/credit lines using active accounts. Total debit must equal total credit.',
'Post the batch. Use the journal list to review status and lines; use Reverse for a posted correction rather than changing the original batch.'
],'Accounts page shows the account structure and balances; Journal Entries retains each batch and its lines. Financial Reports uses the posted journal data.','Financials, Account Mappings, Payroll and Financial Reports.','An account used by journal entries, child accounts or payroll cannot be casually removed. Only posted batches can be reversed; a reversal is a new audit-preserving batch.')
h2('5.4 Financial Reports')
page_guide('Use Financial Reports for accounting statements and balances, after transactions have been cleared/posted.',[
'Open Financial Reports.',
'Choose the statement/report and period controls provided on the page.',
'Review totals and drill to Journal Entries or transaction detail where supporting evidence is required.',
'Use the core reports area for broader transaction detail and profit/loss summaries.'
],'Financial Reports, Core Reports and Journal Entries.','Reports and Financials.')
h2('5.5 Payroll and Payslips')
page_guide('Use Payroll to calculate, submit, approve and pay a monthly employee payroll.',[
'Open Payroll and select New Payroll.',
'Enter payroll name, year, month, pay period, payment date and, when applicable, the payroll expense account.',
'Add each employee line. Selecting a staff member can populate name, position, department and base salary; verify the amounts.',
'Enter bonuses, deductions and tax for each line. The system calculates gross and net salary totals.',
'Save as Draft, then use Submit to send it for review.',
'An authorised approver selects Approve or Reject. Record a rejection reason when rejecting.',
'After approval, select Pay. The payroll moves to Paid and the payroll expense transaction is synchronised to Financials.',
'Open Payslips to view the employee-level results after payroll is created/processed.'
],'Payroll list shows status and totals; Payslips is the employee-level follow-up area. Financials and Journals show the linked financial effect.','Staff, Accounts, Financials, Journal Entries and Payroll Reports.','Only Draft and Rejected payrolls can be edited; only Draft payrolls can be deleted. A payroll period is unique by year and month. Reopen is available for controlled correction before payment.')

h1('6. Construction Management')
h2('6.1 Construction Overview, Projects and Project Detail')
page_guide('Use Construction Overview for portfolio visibility and Projects for creating and managing individual sites.',[
'Open Construction Overview to review project-level cards and current operational indicators.',
'Open Projects and select New construction project.',
'Enter name, description, location, budget, progress, start/end dates, image and initial delivery status.',
'Save. Open the project row to view its detail page.',
'Use the detail page to review linked tasks, assigned staff, daily operational expenses and workforce contracts; use Edit to update the project.'
],'Project list and Project Detail page; associated records remain in their own task, expense, manpower and contract pages.','Staff, Tasks, Project Progress, Expenses, Manpower, Inventory and Reports.','Set a realistic schedule and budget before recording dependent work. Project status and progress should be updated deliberately; they are used for management tracking.')
h2('6.2 Tasks and Project Progress')
page_guide('Use Tasks to assign and track discrete project work, and Project Progress to monitor completion across projects.',[
'Open Project Tasks and select New project task.',
'Choose the project and assigned staff member, then enter task title/details, priority, due date, status and progress.',
'Save. Update the task as ownership, deadline, status or progress changes.',
'Open Project Progress to compare completion and current workload; use task records for the supporting detail.'
],'Project Tasks list, Project Detail and Project Progress.','Projects and Staff.','Keep task progress aligned with task status. Users need task update permission to change assignments, progress or dates.')
h2('6.3 Operational Expenses and Worker Ledger')
page_guide('Use Operational Expenses for site costs and Worker Ledger for labour-related income/expense entries.',[
'Open Operational Expenses and select Record site expense.',
'Enter amount, description, category, expense date and link it to a project and/or staff member when relevant.',
'Save; correct errors using Edit while you have permission.',
'For a labour-specific entry, open Worker Ledger, choose Income or Expense, amount, description, date, project and staff, then save.'
],'Expense list, Worker Ledger, the project detail’s expense section and Financials.','Projects, Staff, Financials and Construction Reports.','These records are synchronised with the unified financial ledger. Use the correct project and staff link so job costing and labour analysis remain accurate.')
h2('6.4 Worker Types, Manpower and Workforce Contracts')
page_guide('Use Worker Types and Manpower to organise the workforce; use Workforce Contracts to control a contract budget and payments.',[
'Create Worker Types first when you need reusable classifications; enter name, display color and description.',
'Open Manpower to review/manage worker information and assignments available in the workspace.',
'Open Workforce Contracts and create a contract linked to a project, with title, description, original budget, dates and notes.',
'Assign workers to the contract and activate it through the available status action.',
'Record a contract payment only for an Active or Completed contract; choose the assigned worker, amount, date, description and notes.',
'Use budget adjustment actions when an authorised change is needed, then complete or close the contract through its controlled status actions.'
],'Contract detail shows assigned workers, payments, adjustments, total paid and remaining budget. Project Detail also lists contracts.','Projects, Staff, Manpower, Financials and Construction Reports.','Contract status changes are controlled. A payment cannot exceed the remaining contract budget and requires the workforce-contract payment permission. Payments create financial records.')
h2('6.5 Construction Inventory')
page_guide('Use Construction Inventory to record materials moving into, out of or used on projects.',[
'Open Construction Inventory and select the action to record a movement.',
'Choose material, movement type, quantity, date and the project where applicable.',
'Use receipt/addition movements for incoming stock and usage/issue movements for stock consumed by a project.',
'Save and review the movement list, balances and project reporting.'
],'Construction Inventory list and construction inventory reports.','Material Products, Projects, Financials and Construction Reports.','Use a project on usage movements so the system can track job costing. Do not record negative/unsupported stock movements; correct an error with the appropriate authorised adjustment process.')

h1('7. Real Estate Management')
h2('7.1 Real Estate Overview and Properties')
page_guide('Use Real Estate Overview for portfolio-level information and Properties to maintain sale/rental listings.',[
'Open Properties and select Add property.',
'Enter title, property type, status, address, description, price, area, bedrooms, bathrooms and optional image.',
'Save. Open the property detail to review all linked deals, rental contracts, tenants and financial transactions.',
'Use Edit to maintain availability, valuation, specifications or imagery.'
],'Properties list and Property Detail page.','Clients, Deals, Rentals, Rent Payments, Financials and Real Estate Reports.','Do not mark a property available if an active rental contract or completed deal makes that inaccurate; review the property detail before changing status.')
h2('7.2 Clients, Deals and Sales')
page_guide('Use Clients to build your buyer/seller/investor directory, Deals to record a sale or rental transaction, and Property Sales for sales-focused tracking.',[
'Create or confirm the client record: name is required; email, phone and notes are optional.',
'Create a Property Deal and select the property and client.',
'Choose SALE or RENTAL, then enter total amount, paid amount, payment status, closed date and notes.',
'Save and keep payment status updated as collection progresses.',
'Open the deal detail to review its linked financial transactions; open Property Sales for sales-oriented review.'
],'Client Detail shows associated deals; Deal Detail and Property Detail show related finance.','Properties, Financials, Property Sales, Reports.','Payment status options include Paid, Partial, Pending, Overdue and Refunded. Use amounts consistently: paid amount should reflect actual money received.')
h2('7.3 Tenants, Rental Contracts and Rent Payments')
page_guide('Use this sequence to manage leasing from tenant registration through collection follow-up.',[
'Create the tenant record with name, contact details, identification/passport number and notes.',
'Create a rental contract by choosing the tenant and property, entering monthly rent, start date, end date, renewal date (if any), status and notes.',
'Use Rental Hub to review the operational rental picture and upcoming/active agreements.',
'Create a rent payment record. Choose the contract to populate the tenant and monthly amount where available; enter due date, paid date, amount due, amount paid, receipt number and notes.',
'Use Mark paid, Mark partial or Mark late to keep the collection status current.',
'Review contract and payment status, then use Financials and real-estate reports for financial follow-up.'
],'Rental Hub, Rental Contracts, Rent Payments, Property Detail and Financials.','Properties, Tenants, Financials and Real Estate Reports.','An active contract needs valid start and end dates. Rent-payment status is Paid, Unpaid, Late or Partial. Link payment to its contract whenever possible to retain the tenant/property context.')

h1('8. Materials Management')
h2('8.1 Materials Overview and Product Inventory')
page_guide('Use Materials Overview for summary visibility and Inventory to maintain stock products and thresholds.',[
'Open Materials Inventory to review stock on hand, product information and low-stock signals.',
'Select Manage material products, then Add material product.',
'Enter the product identity, image, cost, selling price, warehouse/location and stock threshold fields provided.',
'Save and return to inventory to monitor available quantity and exceptions.'
],'Materials Inventory and product-management list.','Suppliers, Purchases, Sales, Transportation and Materials Reports.','Keep the cost, selling price and low-stock threshold current. Product records are used in purchase and sales line items, so verify units and pricing before transacting.')
h2('8.2 Suppliers and Purchase Orders')
page_guide('Use Suppliers to maintain procurement contacts and Purchase Orders to order and receive stock.',[
'Create a supplier with contact details and opening balance information as applicable.',
'Create a purchase order: choose supplier, order/date information and material line items with quantities and costs.',
'Save as Draft, then select Mark ordered when the order is placed.',
'When stock arrives, select Receive stock. The system updates stock, weighted cost and supplier-ledger information.',
'If an order will not proceed, select Cancel; use Reopen draft only for a cancelled order that needs correction and reissue.'
],'Purchase Orders list provides status; inventory and supplier records show downstream effects.','Products, Suppliers, Inventory, Financials and Materials Reports.','Purchase orders cannot be edited from the list once created under the configured workflow. Receive stock only after verifying actual quantities and supplier documentation.')
h2('8.3 Customers and Material Sales')
page_guide('Use Customers to maintain sales accounts and Material Sales to issue invoices and reduce stock.',[
'Create a customer record with the contact information required for billing and follow-up.',
'Create a sales invoice and choose the customer, material line items, quantities, prices, discounts and notes.',
'Review the calculated invoice total and save.',
'Use the printable invoice action when a customer document is required.',
'Review inventory, customer account/receivable information and Financials after the sale.'
],'Material Sales list and printable invoice; Inventory, Customer records and Financials provide follow-up.','Products, Customers, Inventory, Financials and Materials Reports.','Material sales synchronise stock reversal, discounts, revenue and receivable information. Check available stock and invoice amounts before saving.')
h2('8.4 Transportation')
page_guide('Use Transportation to manage a material delivery from dispatch through completion.',[
'Open Transportation and select New delivery record.',
'Enter the transport/delivery details, linked material/order/customer context, cost and notes as the form provides.',
'Save as Pending, then select Start transit at dispatch.',
'Select Mark delivered after delivery is confirmed; the delivery expense is posted when delivered.',
'Use Cancel only when the delivery will not take place.'
],'Transportation list shows Pending, In Transit, Delivered or Cancelled state. Financials shows the delivered expense effect.','Purchases, Sales, Materials Inventory, Customers and Financials.','Do not mark a delivery delivered until it is confirmed; that action has a financial consequence.')

h1('9. Reports, Analytics and Tracking')
h2('9.1 Workspace Reports')
purpose('Each workspace has a Reports entry. Construction, Real Estate and Materials reports support progressively detailed navigation: select a project/property/product-context record where shown, select a category, then open a transaction/item. Core reports include financial and payroll reporting. Use available period/status filters before exporting or acting.')
table(d,['Need','Start here','Follow-up'],[('Financial performance','Financial Reports or Core Reports','Journal Entries and Financials'),('Construction cost/progress','Construction Reports','Project Detail, Tasks, Expenses, Inventory and Contracts'),('Property/revenue/collections','Real Estate Reports','Property, Deal, Contract and Rent Payment detail'),('Stock, procurement and sales','Materials Reports','Inventory, Purchase Orders, Sales and Transportation'),('Payroll results','Payroll Reports and Payslips','Payroll record and Staff')],[2.1,2.35,2.15])
h2('9.2 Report Schedules')
page_guide('Use Report Schedules when an authorised reports administrator needs routine generated reporting.',[
'Open Report Schedules.',
'Create or edit the schedule using the report type, delivery/timing and recipient options provided.',
'Enable the schedule after checking its scope and timing.',
'Return to the schedules list to edit, disable or review the setup.'
],'Report Schedules and the underlying report page.','Core and workspace Reports.','Only users with report-administration permission can manage schedules. Confirm recipients before enabling automated delivery.')

h1('10. Administration and Settings')
h2('10.1 Settings')
page_guide('Use Settings for company-level operational preferences and configuration.',[
'Open Settings.',
'Review the available company, branding, module and preference sections.',
'Update only the section you are responsible for and save.',
'Check the related workspace after a material configuration change.'
],'Settings retains the current company configuration.','All modules.','Settings are shared company controls. Restrict changes to authorised administrators and document policy changes internally.')
h2('10.2 Notifications and Audits')
page_guide('Use Notifications for alerts and Audits for traceability of significant system activity.',[
'Open Notifications from the header or menu and review unread items.',
'Open Audits to search/filter authorised activity records.',
'Use audit information to understand who changed a record and when, then open the owning module to inspect the current record.'
],'Notifications page and Audits page.','All modules.','Audit-log access is limited to authorised leadership roles; do not treat it as a replacement for normal operational reports.')
h2('10.3 Platform administration')
purpose('Super Admin pages (Companies, Company Onboarding, Company details, Billing and Super Admin Account) are for the platform operator, not normal client users. They manage tenant companies, subscriptions/billing and platform accounts. Keep these duties separate from a client company’s day-to-day operations.')

h1('11. Complete Business Workflows')
h2('11.1 Construction delivery and cost control')
table(d,['Stage','Action','Track / next step'],[('Start','Create staff, worker types and a construction project with budget/schedule.','Staff; Projects; Project Detail'),('Plan','Create tasks, assign people, set priority/due dates and status.','Tasks; Project Progress'),('Execute','Record site expenses, worker-ledger entries, inventory usage and contract labour.','Expenses; Worker Ledger; Inventory; Contracts'),('Control','Record contract payments within remaining budget and update task/project progress.','Contract Detail; Project Detail'),('Review','Use construction reports to compare operations and costs.','Construction Reports; Financials')],[1.05,3.5,2.05])
h2('11.2 Property sale')
table(d,['Stage','Action','Track / next step'],[('Start','Create property listing and client record.','Properties; Clients'),('Create deal','Select property/client, set Sale, values, paid amount, status and closing date.','Deals; Deal Detail'),('Collect','Update payment status/amount as money is received.','Deal Detail; Financials'),('Review','Use Property Sales, property detail and real-estate reports.','Property Sales; Reports')],[1.05,3.5,2.05])
h2('11.3 Rental collection')
table(d,['Stage','Action','Track / next step'],[('Start','Create tenant and rental contract with correct property, rent and dates.','Tenants; Rental Contracts'),('Bill/record','Create rent payment linked to contract; set due and paid amounts/dates.','Rent Payments; Rental Hub'),('Update','Mark Paid, Partial or Late as collection changes.','Rent Payments; Financials'),('Follow up','Review contract, property detail and reports for overdue items.','Rental Hub; Property Detail; Reports')],[1.05,3.5,2.05])
h2('11.4 Procure, receive and sell materials')
table(d,['Stage','Action','Track / next step'],[('Prepare','Create product and supplier; set accurate prices/costs and threshold.','Inventory; Suppliers'),('Order','Create draft purchase order then Mark ordered.','Purchase Orders'),('Receive','Verify delivery and Receive stock.','Inventory; Supplier balance; Financials'),('Sell','Create customer then sales invoice with products/discounts.','Sales; printable invoice; Inventory'),('Deliver','Create transport record, Start transit and Mark delivered.','Transportation; Financials; Reports')],[1.05,3.5,2.05])
h2('11.5 Payroll month-end')
table(d,['Stage','Action','Track / next step'],[('Prepare','Confirm active staff, salaries and expense account.','Staff; Accounts'),('Calculate','Create payroll and employee lines; verify bonuses, deductions and tax.','Payroll Draft'),('Approval','Submit; authorised user approves or rejects with reason.','Payroll status'),('Pay','Select Pay only after approval.','Payroll; Payslips; Financials'),('Review','Use payroll reports and financial reports.','Reports; Journals')],[1.05,3.5,2.05])

h1('12. Common Tasks / Quick Guides')
table(d,['Task','Fast path'],[('Add a new employee','Staff > Add staff member > complete required fields > optionally create login > Save.'),('Create a construction task','Construction > Tasks > New > project + owner + deadline + priority/status > Save.'),('Record a site cost','Construction > Expenses > Record site expense > link project/staff > Save.'),('Create a property lease','Real Estate > Rental Contracts > New > tenant + property + rent + dates > Save.'),('Record rent collection','Real Estate > Rent Payments > New > choose contract + amounts/dates > Save > Mark status.'),('Receive purchased stock','Materials > Purchases > open order > Receive stock after checking delivery.'),('Issue a material invoice','Materials > Sales > New > customer + products/quantities/discount > Save/Print.'),('Run payroll','Payroll > New > employee lines > Save > Submit > Approve > Pay.'),('Correct a cleared financial item','Financials > edit authorised transaction; system handles linked journal reversal/re-posting.'),('Change a user role','Staff > select employee > Change role > Save > user signs in again.')],[2.0,4.6])

h1('13. Troubleshooting and Common User Mistakes')
table(d,['Situation','What to check'],[('I cannot see a page or Create button.','Your role/direct permissions may not include the workspace or create action. Ask an administrator to check Roles & Permissions.'),('I cannot edit payroll.','Only Draft and Rejected payrolls can be edited. Use the permitted transition or create a corrective payroll approach.'),('A payroll cannot be paid.','It must be Approved first; confirm the approver has completed the transition.'),('A contract payment is blocked.','Confirm the contract is Active or Completed, the worker is assigned and the payment does not exceed remaining budget.'),('I cannot receive stock.','Verify the purchase order is Draft or Ordered, and that the physical delivery is confirmed.'),('A delivery should not be marked delivered yet.','Keep it Pending/In Transit; delivered status posts the delivery expense.'),('A rent payment has the wrong context.','Edit/link the rental contract so tenant and monthly-rent information remains consistent.'),('An accounting batch needs correction.','Reverse the posted batch; do not attempt to alter the original posted record.'),('An account cannot be deleted.','It is likely used by journals, child accounts or payroll. Keep/archive it according to company policy.'),('A user cannot sign in after access changes.','Role/permission updates require a new sign-in. Confirm the staff account is active and password meets policy.')],[2.15,4.45])

h1('14. Best Practices')
add_bullets(d,['Create master records before transactions: staff before payroll/tasks, properties before deals/contracts, products before purchases/sales.','Use statuses immediately and accurately; status controls which next actions are available and makes reports useful.','Always link transactions to their project, property, contract, deal, supplier or customer whenever the form allows it.','Use clear descriptions and receipt/reference numbers so another authorised user can audit the record later.','Review Dashboard, workspace reports and Financial Reports routinely; investigate unexpected figures from the underlying list rather than changing totals blindly.','Separate preparation, approval and payment duties where possible through roles and permissions.','Use reversal and controlled status actions for corrections; avoid duplicating records to “fix” an error.','Review low-stock, overdue rent, pending approvals and upcoming contract/lease dates on a regular schedule.'])

h1('15. Glossary of Important Terms')
table(d,['Term','Meaning'],[('Account','A chart-of-accounts category used to classify accounting entries.'),('Cleared transaction','A financial transaction ready for accounting posting when configuration permits.'),('Deal','A property sale or rental transaction involving a property and client.'),('Journal batch','A balanced accounting record containing debit and credit lines.'),('Payroll period','A month/year payroll record containing calculated employee items.'),('Property status','The operational availability state of a property listing.'),('Rental contract','A lease that links a tenant, property, rent amount and dates.'),('Status transition','A controlled change such as Submit, Approve, Pay, Receive stock or Mark delivered.'),('Workforce contract','A project labour contract with budget, worker assignments, payments and adjustments.'),('Weighted cost','The inventory cost recalculated when stock is received.'),('Workspace','A functional area of MmaamulPro: Core, Construction, Real Estate or Materials.')],[1.85,4.75])
callout(d,'Coverage check','This manual covers the client-facing application routes, major pages, their input/actions, status lifecycles and cross-module tracking paths. Platform-only Super Admin functions are identified separately to avoid assigning them to client operators.','EAF2F8')

d.save(OUT)
print(OUT)
